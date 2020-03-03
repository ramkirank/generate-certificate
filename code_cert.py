# -*- coding: utf-8 -*-

import docx # if docx is not installed use pip install python-docx    for python3
import pandas

dox = docx.Document('sample_Certificate of participation.docx')
excel_data_df = pandas.read_excel('sample-STUDENT LIST.xlsx')
paras= []
for i in range(len(dox.tables[0].cell(0,0).paragraphs)):
    paras.append(dox.tables[0].cell(0,0).paragraphs[i].text)
    print(i,"--->",paras[i]) 
    

para_number = 3     # "Name" is in paragraph 3 as seen in output
runs = []
for i in range(len(dox.tables[0].cell(0,0).paragraphs[para_number].runs)):
    runs.append(dox.tables[0].cell(0,0).paragraphs[para_number].runs[i].text)
    print(i,"--->",runs[i])
run_number = 3      #  "Name " is in index 3 in runs

coloum_number = 0   # enter coloum number of "Names" list in excel
for i in range(excel_data_df.shape[coloum_number]):
    dox.tables[0].cell(0,0).paragraphs[para_number].runs[run_number].text = excel_data_df['Student Name'][i]
    dox.save('Test_certificate'+str(i+1)+'.docx')

 
    
    